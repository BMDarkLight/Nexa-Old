import logging
from pydantic import BaseModel
from typing import List, Dict, Any

from fastapi import HTTPException

import traceback

import PyPDF2, io
import docx
import pandas as pd

from api.schemas.base import PyObjectId

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_content: bytes) -> str:
    logger.info("Starting extraction of text from PDF file.")
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        num_pages = len(pdf_reader.pages)
        logger.debug(f"Number of pages in PDF: {num_pages}")
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        logger.info(f"Successfully extracted text from PDF file. Extracted text length: {len(text)} characters, Pages: {num_pages}")
        return text
    except Exception as e:
        logger.error(f"Failed to extract text from PDF file: {e}", exc_info=True)
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF file: {e}")

def extract_text_from_docx(file_content: bytes) -> str:
    logger.info("Starting extraction of text from DOCX file.")
    try:
        doc = docx.Document(io.BytesIO(file_content))
        num_paragraphs = len(doc.paragraphs)
        logger.debug(f"Number of paragraphs in DOCX: {num_paragraphs}")
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        logger.info(f"Successfully extracted text from DOCX file. Extracted text length: {len(text)} characters, Paragraphs: {num_paragraphs}")
        return text
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX file: {e}", exc_info=True)
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX file: {e}")

def _load_spreadsheet(file_content: bytes, file_type: str) -> pd.DataFrame:
    if file_type == 'excel':
        return pd.read_excel(io.BytesIO(file_content), engine='openpyxl')
    elif file_type == 'csv':
        try:
            return pd.read_csv(io.BytesIO(file_content))
        except UnicodeDecodeError:
            logger.warning("CSV not UTF-8 encoded; using latin1 fallback.")
            return pd.read_csv(io.BytesIO(file_content), encoding='latin1')
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def extract_text_from_excel(file_content: bytes) -> str:
    logger.info("Starting extraction of text from Excel file.")
    try:
        df = _load_spreadsheet(file_content, 'excel')
        num_columns = len(df.columns)
        num_rows = len(df)
        logger.debug(f"Excel columns: {list(df.columns)}")
        text_lines = []

        # Add headers
        headers = [str(col).strip() for col in df.columns]
        text_lines.append("\t".join(headers))

        # Add first 5 rows or fewer
        sample_rows = df.head(5)
        for _, row in sample_rows.iterrows():
            row_values = [str(val).strip() if pd.notnull(val) else '' for val in row]
            text_lines.append("\t".join(row_values))

        text = "\n".join(text_lines) + "\n"
        logger.info(f"Successfully extracted text from Excel file. Extracted text length: {len(text)} characters, Columns: {num_columns}, Rows: {num_rows}")
        return text
    except Exception as e:
        logger.error(f"Failed to extract text from Excel file: {e}", exc_info=True)
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=f"Failed to extract text from Excel file: {e}")

def extract_text_from_csv(file_content: bytes) -> str:
    logger.info("Starting extraction of text from CSV file.")
    try:
        df = _load_spreadsheet(file_content, 'csv')
        num_columns = len(df.columns)
        num_rows = len(df)
        logger.debug(f"CSV columns: {list(df.columns)}")
        text_lines = []

        # Add headers
        headers = [str(col).strip() for col in df.columns]
        text_lines.append("\t".join(headers))

        # Add first 5 rows or fewer
        sample_rows = df.head(5)
        for _, row in sample_rows.iterrows():
            row_values = [str(val).strip() if pd.notnull(val) else '' for val in row]
            text_lines.append("\t".join(row_values))

        text = "\n".join(text_lines) + "\n"
        logger.info(f"Successfully extracted text from CSV file. Extracted text length: {len(text)} characters, Columns: {num_columns}, Rows: {num_rows}")
        return text
    except Exception as e:
        logger.error(f"Failed to extract text from CSV file: {e}", exc_info=True)
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=f"Failed to extract text from CSV file: {e}")

def extract_table_from_excel(file_content: bytes) -> Dict[str, Any]:
    logger.info("Starting extraction of table from Excel file.")
    try:
        df = _load_spreadsheet(file_content, 'excel')
        num_columns = len(df.columns)
        num_rows = len(df)
        logger.debug(f"Excel DataFrame columns: {list(df.columns)}, shape: {df.shape}")
        schema = {col: str(dtype) for col, dtype in df.dtypes.items()}
        sample = df.head(5).to_dict(orient='records')
        shape = df.shape
        logger.info(f"Successfully extracted table from Excel file. Columns: {num_columns}, Rows: {num_rows}")
        return {
            "schema": schema,
            "sample": sample,
            "shape": shape,
            "dataframe": df
        }
    except Exception as e:
        logger.error(f"Failed to extract table from Excel file: {e}", exc_info=True)
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=f"Failed to extract table from Excel file: {e}")

def extract_table_from_csv(file_content: bytes) -> Dict[str, Any]:
    logger.info("Starting extraction of table from CSV file.")
    try:
        df = _load_spreadsheet(file_content, 'csv')
        num_columns = len(df.columns)
        num_rows = len(df)
        logger.debug(f"CSV DataFrame columns: {list(df.columns)}, shape: {df.shape}")
        schema = {col: str(dtype) for col, dtype in df.dtypes.items()}
        sample = df.head(5).to_dict(orient='records')
        shape = df.shape
        logger.info(f"Successfully extracted table from CSV file. Columns: {num_columns}, Rows: {num_rows}")
        return {
            "schema": schema,
            "sample": sample,
            "shape": shape,
            "dataframe": df
        }
    except Exception as e:
        logger.error(f"Failed to extract table from CSV file: {e}", exc_info=True)
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=f"Failed to extract table from CSV file: {e}")